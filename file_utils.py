from docx import Document
import tempfile

def load_text(uploaded_file):
    if uploaded_file.name.endswith(".docx"):
        doc = Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text, "docx"
    else:
        text = uploaded_file.read()
        return text, "txt"

def convert_text_to_docx(text):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_path.name)
    return temp_path.name

def convert_text_to_txt(text):
    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode='w', encoding='utf-8')
    temp_path.write(text)
    temp_path.close()
    return temp_path.name
